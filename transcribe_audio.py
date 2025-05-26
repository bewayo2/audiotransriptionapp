import openai
import os
from dotenv import load_dotenv
from docx import Document
import re

# Load environment variables from .env file
load_dotenv()

# Get variables from environment
openai.api_key = os.getenv("OPENAI_API_KEY")
audio_path = os.getenv("TRANSCRIPTION_AUDIO_PATH")
output_dir = os.getenv("TRANSCRIPTION_OUTPUT_DIR")

if not openai.api_key or not audio_path or not output_dir:
    raise ValueError("OPENAI_API_KEY, TRANSCRIPTION_AUDIO_PATH, and TRANSCRIPTION_OUTPUT_DIR must be set in the .env file.")

os.makedirs(output_dir, exist_ok=True)
output_basename = os.path.splitext(os.path.basename(audio_path))[0]
output_txt_path = os.path.join(output_dir, output_basename + ".txt")
output_docx_path = os.path.join(output_dir, output_basename + ".docx")

# Transcribe the audio
with open(audio_path, "rb") as audio_file:
    transcript = openai.audio.transcriptions.create(
        model="whisper-1",
        file=audio_file
    )

# Save the transcription to a text file
with open(output_txt_path, "w", encoding="utf-8") as f:
    f.write(transcript.text)

# Save the transcription to a formatted Word document
# Split into paragraphs at sentence ends for readability
paragraphs = re.split(r'(?<=[.!?]) +', transcript.text)
doc = Document()
heading_text = f'Transcription for "{os.path.splitext(os.path.basename(audio_path))[0]}"'
doc.add_heading(heading_text, 0)
for para in paragraphs:
    doc.add_paragraph(para.strip())
doc.save(output_docx_path)

print(f"Transcription saved to: {output_txt_path}")
print(f"Formatted transcription saved to: {output_docx_path}") 